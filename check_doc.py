from docx import Document

doc = Document(r'C:\Users\Ayant\Desktop\DREAM_CAREER\report\DreamCareer_Report.docx')
headings = [(p.style.name, p.text[:100]) for p in doc.paragraphs if p.style.name.startswith('Heading')]
for h in headings:
    print(h[0], ':', h[1])
print('Total headings:', len([p for p in doc.paragraphs if p.style.name.startswith('Heading')]))