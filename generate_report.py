#!/usr/bin/env python
"""
Generate DreamCareer Report as DOCX using python-docx
Converts all markdown files in report/ to a single DOCX with proper formatting
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import sys

# Add project root to path
PROJECT_ROOT = r"C:\Users\Ayant\Desktop\DREAM_CAREER"
REPORT_DIR = os.path.join(PROJECT_ROOT, "report")

# Chapter files in order
CHAPTER_FILES = [
    "Ch1_Introduction.md",
    "Ch2_Literature_Review.md",
    "Ch3_Planning.md",
    "Ch4_Requirements.md",
    "Ch5_System_Design.md",
    "Ch6_Implementation.md",
    "Ch7_Testing.md",
    "Ch8_Conclusion.md",
    "Appendix_Questionnaire.md",
    "Appendix_Gantt_Chart.md",
]

def create_document():
    """Create the main document with styles"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        if level == 1:
            heading_style.font.size = Pt(22)
            heading_style.font.bold = True
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        elif level == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.bold = True
    
    # Configure code style
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Cm(1)
    
    # Code block style
    code_block_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_block_style.font.name = 'Consolas'
    code_block_style.font.size = Pt(8.5)
    code_block_style.paragraph_format.space_before = Pt(6)
    code_block_style.paragraph_format.space_after = Pt(6)
    code_block_style.paragraph_format.left_indent = Cm(1.5)
    code_block_style.paragraph_format.right_indent = Cm(1.5)
    
    return doc

def parse_markdown_line(line):
    """Parse a markdown line and return (style, text, is_code_block, is_list)"""
    line = line.rstrip()
    
    # Code block detection
    if line.startswith('```'):
        return ('code_block', line[3:].strip(), True, False)
    
    # Headers
    if line.startswith('### '):
        return ('Heading 3', line[4:].strip(), False, False)
    elif line.startswith('## '):
        return ('Heading 2', line[3:].strip(), False, False)
    elif line.startswith('# '):
        return ('Heading 1', line[2:].strip(), False, False)
    
    # List items
    if line.startswith('- ') or line.startswith('* '):
        return ('List Bullet', line[2:].strip(), False, True)
    if re.match(r'^\d+\.\s', line):
        return ('List Number', re.sub(r'^\d+\.\s*', '', line), False, True)
    
    # Bold/italic inline - we'll handle in paragraph processing
    # Code inline
    if line.startswith('`') and line.endswith('`') and len(line) > 2:
        return ('Code', line[1:-1], False, False)
    
    # Horizontal rule
    if line.strip() in ('---', '***', '***'):
        return ('HorizontalRule', '', False, False)
    
    # Table row
    if '|' in line and line.count('|') >= 2:
        return ('TableRow', line.strip(), False, False)
    
    # Empty line
    if not line.strip():
        return ('Empty', '', False, False)
    
    return ('Normal', line, False, False)

def process_table_row(doc, line):
    """Process a markdown table row"""
    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
    if not cells:
        return
    
    table = None
    # Find the last table in the document
    for element in doc.element.body:
        if element.tag.endswith('}tbl'):
            table = element
            break
    
    if table is None:
        # Create new table - need to convert CT_Tbl to Table object
        from docx.table import Table
        table = Table(doc.element.body[-1], doc) if doc.element.body else doc.add_table(rows=1, cols=len(cells))
        if hasattr(table, 'style'):
            table.style = 'Table Grid'
        else:
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
        # Header row
        for i, cell_text in enumerate(cells):
            if i < len(table.rows[0].cells):
                cell = table.rows[0].cells[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
    else:
        # Add row to existing table - need to convert to Table object
        from docx.table import Table
        try:
            table_obj = Table(table, doc)
            row = table_obj.add_row()
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
        except:
            # Fallback: create new table
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            for i, cell_text in enumerate(cells):
                if i < len(table.rows[0].cells):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)

def add_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with bold/italic processing"""
    p = doc.add_paragraph(style=style)
    
    # Process **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
    return p

def process_markdown_file(doc, filepath):
    """Process a markdown file and add to document"""
    if not os.path.exists(filepath):
        print(f"  WARNING: File not found: {filepath}")
        return
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    in_table = False
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(style='CodeBlock')
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                in_code_block = False
                code_lines = []
            else:
                # Start code block
                in_code_block = True
                code_lines = []
                lang = line[3:].strip()
                if lang:
                    p = doc.add_paragraph(style='Code')
                    run = p.add_run(f"Language: {lang}")
                    run.italic = True
                    run.font.size = Pt(8)
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Parse the line
        style, text, is_code, is_list = parse_markdown_line(line)
        
        if style == 'Heading 1':
            doc.add_heading(text, level=1)
        elif style == 'Heading 2':
            doc.add_heading(text, level=2)
        elif style == 'Heading 3':
            doc.add_heading(text, level=3)
        elif style == 'Code':
            p = doc.add_paragraph(style='Code')
            run = p.add_run(text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif style == 'List Bullet':
            doc.add_paragraph(text, style='List Bullet')
        elif style == 'List Number':
            doc.add_paragraph(text, style='List Number')
        elif style == 'TableRow':
            process_table_row(doc, text)
        elif style == 'HorizontalRule':
            doc.add_paragraph('—' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == 'Empty':
            continue  # Skip empty lines
        else:
            # Normal paragraph with inline formatting
            p = add_formatted_paragraph(doc, text)

def main():
    """Main function to generate the report"""
    print("Generating DreamCareer Report...")
    
    doc = create_document()
    
    # Title page
    doc.add_heading('DreamCareer', level=0)
    p = doc.add_paragraph()
    run = p.add_run('AI-Powered Career Intelligence Platform for Sri Lanka')
    run.font.size = Pt(16)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('BSc (Hons) Data Science — Final Year Project').font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DreamCareer: AI-Powered Career Intelligence Platform')
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Final Year Project Report')
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BSc (Hons) Data Science')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026')
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph('(Generated automatically in Word — right-click → Update Field)')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()
    
    # Process each chapter
    chapter_files = [
        ("Ch1_Introduction.md", "Chapter 1: Introduction"),
        ("Ch2_Literature_Review.md", "Chapter 2: Literature Review"),
        ("Ch3_Planning.md", "Chapter 3: Planning"),
        ("Ch4_Requirements.md", "Chapter 4: Requirements Gathering and Analysis"),
        ("Ch5_System_Design.md", "Chapter 5: System Design"),
        ("Ch6_Implementation.md", "Chapter 6: Implementation"),
        ("Ch7_Testing.md", "Chapter 7: Testing and Validation"),
        ("Ch8_Conclusion.md", "Chapter 8: Conclusion"),
        ("Appendix_Questionnaire.md", "Appendix A: Questionnaire and Interview Data"),
        ("Appendix_Gantt_Chart.md", "Appendix B: Gantt Chart"),
    ]
    
    for filename, title in chapter_files:
        filepath = os.path.join("report", filename)
        if os.path.exists(filepath):
            print(f"Processing {filename}...")
            doc.add_heading(title, level=1)
            process_markdown_file(doc, os.path.join("report", filename))
            doc.add_page_break()
        else:
            print(f"  WARNING: {filename} not found")
    
    # References
    doc.add_heading('References', level=1)
    p = doc.add_paragraph('(Generated from references.bib — compile with pandoc --citeproc for full bibliography)')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    # Save
    output_path = os.path.join("report", "DreamCareer_Report.docx")
    doc.save(output_path)
    print(f"\n✅ Report saved to: {output_path}")

if __name__ == '__main__':
    main()